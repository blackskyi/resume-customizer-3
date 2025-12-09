from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from anthropic import Anthropic

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
app.config['OUTPUT_FOLDER'] = '/tmp/outputs'
app.config['TEMPLATES_FOLDER'] = os.path.join(os.path.dirname(__file__), 'resume_templates')

# Create folders if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
os.makedirs(app.config['TEMPLATES_FOLDER'], exist_ok=True)

# Resume content for different roles (will be customized by Claude)
RESUME_TEMPLATES = {
    'verification': 'verification_fpga_rtl',
    'devops': 'devops',
    'physical_design': 'physical_design',
    'dft': 'dft'
}

# Initialize Claude API client
claude_client = None
api_key = os.environ.get('ANTHROPIC_API_KEY')
if api_key:
    try:
        claude_client = Anthropic(api_key=api_key)
        print('✅ Claude API initialized successfully')
    except Exception as e:
        print(f'⚠️ Claude API initialization failed: {e}')
else:
    print('⚠️ ANTHROPIC_API_KEY not set')

def get_base_resume_content(role_type):
    """Get the base resume content for a specific role"""
    
    if role_type == 'devops':
        return """
GOKUL PK
Email: gprasann@usc.edu | Phone: +1 213 316 8527
LinkedIn: linkedin.com/in/gokul-pk-4b16b9345

PROFESSIONAL SUMMARY
Senior DevOps Engineer with 7+ years driving cloud infrastructure automation and CI/CD excellence. Reduced deployment time by 60% through GitOps implementation and cut infrastructure costs by 35% via containerization strategies. Expert in building scalable Kubernetes platforms, implementing DevSecOps pipelines, and architecting service mesh solutions that support 99.9% uptime for enterprise applications.

TECHNICAL SKILLS
• Cloud & Infrastructure: AWS (EC2, S3, EKS, Lambda, CloudFormation), Azure (AKS, ARM Templates), GCP
• Container Orchestration: Kubernetes, Docker, Helm, ArgoCD, OpenShift, Kustomize
• CI/CD & Automation: Jenkins, GitHub Actions, GitLab CI/CD, Tekton, Bamboo, TeamCity
• Infrastructure as Code: Terraform, Pulumi, Ansible, Chef, CloudFormation, ARM Templates
• Service Mesh & Observability: Istio, Linkerd, Prometheus, Grafana, Datadog, OpenTelemetry, Jaeger
• Security & Compliance: SonarQube, Trivy, DevSecOps practices, IAM, Security Groups
• Languages & Scripting: Python, Bash, PowerShell, Go, Groovy, Ruby, YAML, JSON, HCL
• Databases: PostgreSQL, MySQL, MongoDB, RDS, DynamoDB

WORK EXPERIENCE

Senior DevOps Engineer | Silicon Bricks | Aug 2018 - Dec 2023
• Architected GitOps-based continuous deployment using ArgoCD and Kubernetes, reducing deployment time from 2 hours to 20 minutes (83% improvement) and achieving 99.9% uptime across 50+ microservices
• Led platform engineering reducing infrastructure costs by 35% through containerization with Docker and Kubernetes, optimizing resource utilization from 45% to 78%
• Built cloud-native CI/CD pipelines using Tekton and GitHub Actions with integrated security scanning (Trivy, SonarQube), enabling 200+ weekly deployments with zero critical vulnerabilities
• Implemented Istio service mesh for 40+ microservices, achieving 99.95% request success rate and reducing MTTD by 65% through distributed tracing with OpenTelemetry and Jaeger
• Automated infrastructure provisioning across AWS, Azure, and GCP using Terraform and Pulumi, reducing setup time from 3 days to 4 hours with 100% reproducibility
• Designed DevSecOps practices reducing security rollbacks by 78% and achieving SOC 2 compliance
• Developed reusable Helm charts and Kustomize configurations for Kubernetes deployments, standardizing processes across 15+ teams and reducing configuration errors by 90%
• Created comprehensive monitoring with Prometheus, Grafana, and Datadog, reducing incident response time by 55%
• Mentored team of 6 junior DevOps engineers on cloud-native best practices, improving team velocity by 40%
• Migrated legacy monolithic applications to containerized microservices, improving deployment frequency from monthly to daily releases

DevOps Engineer - CI/CD & Configuration Management | Sanemi Technologies | Jul 2016 - Aug 2018
• Engineered end-to-end Jenkins CI/CD pipelines for 30+ Java and .NET applications, reducing build/deployment time by 70%
• Containerized 25+ legacy applications using Docker and deployed to Kubernetes clusters on Azure AKS, improving resource efficiency by 45%
• Automated infrastructure configuration using Ansible playbooks for 100+ servers, reducing manual errors by 95%
• Developed Chef cookbooks in Ruby for application deployment, managing 200+ EC2 instances across environments
• Implemented Azure migration strategy using Azure Site Recovery, migrating 50+ workloads with zero data loss
• Built custom Docker images and established registry workflows, reducing build time by 60%
• Created automated backup and disaster recovery procedures achieving RPO of 4 hours and RTO of 2 hours
• Designed network security architecture using Azure VNets, security groups, and load balancers, maintaining 99.8% availability

EDUCATION
University of Southern California — Master of Science, Electrical Engineering | Los Angeles, CA
R.V. College of Engineering — Bachelor's in Electronics & Communications Engineering | Bangalore, India
"""
    
    elif role_type == 'verification':
        return """
GOKUL P KUMAR
Email: gocoolpkumar@gmail.com | Phone: +1 213 301 9692
LinkedIn: linkedin.com/in/gokul-p-kumar-23912873

EDUCATION
University of Southern California - Master of Science, Electrical Engineering | Los Angeles, USA
R.V. College of Engineering - Bachelors in Electronics and Communications Engineering | Bangalore, India

PROJECTS
• 32-bit Out-of-Order CPU Design (VHDL, Verilog, SystemVerilog) - Developed micro-architecture implementation for high-performance CPU with 48 physical registers, achieving 20% performance uplift. Integrated copy-free checkpointing (CFC) into SoC RTL, improving speculative execution accuracy by 15%. Ensured functional correctness with store buffer (SB) and store address buffer (SAB), reducing memory disambiguation latency by 10%. Achieved timing closure.

• GPU Compute Tile Design (Assembly Language, Verilog, SystemVerilog) - Engineered GPGPU compute tile with SIMT stack and CUDA cores, enabling 2x faster matrix operations. Performed debugging features, optimized memory coalescing, reduced access latency by 30%.

• Chip Multi-Processor (CMP) Design (Assembly Language, Verilog) - Architected 4-core CPU with MOESI cache coherence, achieving 40% speedup in parallel workloads. Integrated functional blocks into SoC RTL, implementing lockless LL/SC mechanisms.

• Cache Design for Divider (VHDL) - Developed 16x8 fully associative cache CAM with LRU, improving performance by 35%. Performed design quality checks and met timing constraints.

• PCIe PHY Layer Design (Verilog, SystemVerilog) - Engineered PCIe PHY layer with elastic buffer, de-skew FIFO, and 8b10b decoder, achieving 99.9% data integrity across clock crossings. Debugged with Vivado ChipScope.

• AXI Interconnect for SoC (Verilog, SystemVerilog) - Designed 2x4 mesh AXI interconnect, improving data throughput by 28%. Optimized write response buffers, reducing latency by 15%.

KEY SKILLS
• Coding Languages: Python, C, C++, Verilog, VHDL, SystemVerilog, Assembly, Tcl, Perl, Cadence SKILL
• EDA Tools: Cadence Virtuoso, HSPICE, QuestaSim, ModelSim, Vivado, Tetramax, IC-Manage, Synopsys VC Formal, SpyGlass CDC, Verdi
• Design Expertise: ASIC Design, Custom Circuit Design, VLSI, CMOS Technology, RTL coding, RTL-to-GDSII flow, Micro-architecture, SoC integration, CDC, STA, PDK installation
• Verification: UVM, Testbench development, Code Coverage, Functional Coverage, Assertion-based Verification, Formal Property Checking, Fault Simulation, DFT, ATPG, Scan chains, LEC
• Additional Knowledge: RISC-V, ARM, FPGA, DDR, AXI, PCIe, Branch Prediction, Virtual Memory, NoC, Power/Timing Optimization, Neuromorphic computing

EXPERIENCE

Siliconbricks - Design Engineer | Aug 2018 - Dec 2023
• Developed comprehensive UVM testbenches in SystemVerilog achieving 100% functional coverage for complex IP/SoC verification across multiple projects, reducing verification time by 20%
• Designed and implemented UVM verification components including drivers, monitors, sequencers, scoreboards, and agents for industry-standard protocols (AXI, PCIe, DDR)
• Led functional verification efforts for CPU cores, GPU compute tiles, and memory subsystems, identifying and resolving 150+ RTL bugs through systematic simulation and formal property checking
• Achieved 100% code coverage and functional coverage across all verification projects, implementing coverage models and covergroups for comprehensive verification
• Performed clock domain crossing (CDC) verification using Synopsys VC Formal and SpyGlass CDC, identifying and resolving metastability issues
• Automated regression testing and CI/CD flows using Python and Perl scripts integrated with Jenkins, reducing regression runtime by 35%
• Conducted low power verification with UPF-based power-aware simulation, validating power management units (PMU), power gating, and voltage scaling
• Debugged complex verification failures using Synopsys Verdi and waveform analysis, reducing MTTR by 40%

University of Southern California - Teaching Assistant for DFT | Los Angeles, CA | Aug 2024 - Present
• Guided students on fault modeling, test pattern generation, ATPG, and fault simulation
• Supervised hands-on projects including ATPG and Fault Simulation for SoC using C++, Python, and Verilog

University of Southern California - Research Assistant | Los Angeles, CA | Jan 2025 - Present  
• Characterized HfO2 memristors for neuromorphic computing, achieving 25% improvement in multi-level resistive switching accuracy
• Designed and analyzed 12-bit Multiplier and Accumulator (MAC) and 512-bit SRAM in 45nm CMOS, achieving 15% power reduction and 20% faster computation

LEADERSHIP AND ACHIEVEMENTS
• Part of Team Vyoma, which won the NASA Systems Engineering Award
• Multiple Best Short Film Awards showcasing communication skills and creativity
"""
    
    else:
        # Default generic resume
        return """GOKUL PK - Professional Resume"""

def customize_resume_with_claude(base_resume, job_description, role_type):
    """Use Claude to customize resume based on job description"""
    
    if not claude_client:
        print('⚠️ Claude API not available, returning base resume')
        return base_resume
    
    try:
        prompt = f"""You are an expert resume writer and career coach. You will customize a resume to perfectly match a specific job description.

CURRENT RESUME:
{base_resume}

JOB DESCRIPTION:
{job_description}

INSTRUCTIONS:
1. Analyze the job description and identify the key required skills, tools, technologies, and qualifications
2. Reorganize and rewrite the TECHNICAL SKILLS section to emphasize the most relevant skills first
3. Add any missing relevant skills from the job description that the candidate would reasonably have based on their experience
4. Reorder bullet points in the EXPERIENCE section to highlight the most relevant achievements first
5. Keep all the same jobs and projects - DO NOT remove anything, just reorder for relevance
6. Maintain the same professional tone and writing style
7. Keep all quantified achievements (percentages, numbers)
8. DO NOT fabricate experience or achievements
9. Keep the resume to a reasonable length
10. Make it ATS-friendly by using keywords from the job description

Return the COMPLETE customized resume in plain text format. Maintain the same structure and formatting.
"""

        response = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            temperature=0.3,
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )

        customized_resume = response.content[0].text.strip()
        print(f'✅ Resume customized with Claude for {role_type} position')
        return customized_resume

    except Exception as e:
        print(f'⚠️ Error customizing with Claude: {e}')
        return base_resume

def create_docx_from_text(resume_text, output_path):
    """Create a DOCX file from plain text resume"""
    
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    lines = resume_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line
            doc.add_paragraph()
            continue
        
        # Detect headers (all caps or starts with specific keywords)
        if line.isupper() or line.startswith('GOKUL'):
            # Header
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if 'GOKUL' in line or '@' in line else WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.bold = True
            run.font.size = Pt(14) if 'GOKUL' in line else Pt(12)
            run.font.color.rgb = RGBColor(0, 56, 168)  # Blue color
        
        elif line.startswith('•'):
            # Bullet point
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            run = p.runs[0]
            run.font.size = Pt(10)
        
        elif '|' in line and any(keyword in line for keyword in ['Engineer', 'Assistant', 'Technologies', 'Bricks']):
            # Job title line
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.bold = True
            run.font.size = Pt(11)
        
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(10)
    
    doc.save(output_path)
    return output_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/enhance', methods=['POST'])
def enhance():
    job_description = request.form.get('job_description', '').strip()
    role_type = request.form.get('role_type', 'devops')

    if not job_description:
        flash('Please provide a job description')
        return redirect(url_for('index'))

    try:
        # Get base resume content
        base_resume = get_base_resume_content(role_type)
        
        # Customize with Claude
        customized_resume = customize_resume_with_claude(base_resume, job_description, role_type)
        
        # Create DOCX
        output_filename = f'Gokul_PK_{role_type.title()}_Resume_Customized.docx'
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        create_docx_from_text(customized_resume, output_path)
        
        return render_template('result.html',
                               output_filename=output_filename,
                               missing_tools={},
                               gap_report=f"Resume customized for {role_type} position using AI")
    except Exception as e:
        flash(f'Error processing resume: {str(e)}')
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download(filename):
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if os.path.exists(output_path):
        return send_file(output_path, as_attachment=True)
    else:
        flash('File not found')
        return redirect(url_for('index'))

# ============================================================================
# API ENDPOINTS FOR GOOGLE APPS SCRIPT INTEGRATION
# ============================================================================

@app.route('/api/customize-resume', methods=['POST'])
def api_customize_resume():
    """
    API endpoint for Google Apps Script to customize resume
    Expected: form data with 'requirements' text and 'role_type'
    Returns: JSON with success status and filename
    """
    try:
        job_requirements = request.form.get('requirements', '').strip()
        role_type = request.form.get('role_type', 'devops')

        if not job_requirements:
            return {'success': False, 'message': 'No job requirements provided'}, 400

        # Get base resume content
        base_resume = get_base_resume_content(role_type)
        
        # Customize with Claude
        customized_resume = customize_resume_with_claude(base_resume, job_requirements, role_type)
        
        # Create DOCX
        output_filename = f'Gokul_PK_{role_type.title()}_Resume_Customized.docx'
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        create_docx_from_text(customized_resume, output_path)

        # Return success with file path
        return {
            'success': True,
            'message': 'Resume customized successfully',
            'filename': output_filename
        }, 200

    except Exception as e:
        return {'success': False, 'message': f'Error: {str(e)}'}, 500


@app.route('/api/download-resume/<filename>', methods=['GET'])
def api_download_resume(filename):
    """
    API endpoint to download a customized resume
    Returns the file directly
    """
    try:
        # Sanitize filename for security
        safe_filename = secure_filename(filename)
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], safe_filename)
        
        if os.path.exists(output_path):
            return send_file(
                output_path,
                as_attachment=True,
                download_name=safe_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return {'success': False, 'message': 'File not found'}, 404
            
    except Exception as e:
        return {'success': False, 'message': f'Error: {str(e)}'}, 500


@app.route('/api/health', methods=['GET'])
def api_health():
    """
    Health check endpoint
    """
    return {
        'status': 'healthy',
        'service': 'resume-customizer-3-ai',
        'claude_api': 'enabled' if claude_client else 'disabled',
        'version': '3.0-ai-customization'
    }, 200

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
